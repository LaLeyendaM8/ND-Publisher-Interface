from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openai import OpenAI

from app.core.chunking import chunk_pages, slice_pages, truncate_text_if_needed
from app.core.config import require_openai_api_key
from app.core.files import ensure_dir, read_pdf_pages
from app.core.retry import call_with_retries
from app.core.schemas import Artifact, JobResult


SYSTEM_PROMPT = (
    "Actuas como editor bibliografico y traductor especializado en referencias academicas aleman-espanol. "
    "Tu tarea no es traducir prosa normal, sino procesar bibliografia y notas con intervencion minima."
)


@dataclass
class BibliographyConfig:
    pages_per_chunk: int = 1
    model: str = "gpt-4.1"
    temperature: float = 0.1
    timeout_seconds: float = 180.0
    retries: int = 5
    retry_base_sleep: float = 2.0
    max_chars_per_chunk: int = 14000
    add_chunk_headers: bool = True
    page_start: int = 0
    page_end: int | None = None


def _append_jsonl(path: Path, row: dict[str, Any]) -> None:
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def _build_docx_from_jsonl(jsonl_path: Path, out_docx: Path, add_headers: bool) -> None:
    doc = Document()
    with jsonl_path.open("r", encoding="utf-8") as handle:
        for line in handle:
            obj = json.loads(line)
            if add_headers:
                pages = ", ".join(str(p + 1) for p in obj["page_indices"])
                doc.add_heading(f"Chunk {obj['chunk_id']} (PDF Seiten: {pages})", level=2)
            doc.add_paragraph(obj["translation"])
    doc.save(str(out_docx))


def _process_chunk(client: OpenAI, cfg: BibliographyConfig, text: str) -> str:
    prompt = f"""
Procesa la siguiente pagina o bloque de notas/bibliografia.

REGLAS ESTRICTAS:
1. Conserva el texto original al maximo.
2. Manten numeracion y orden exacto.
3. No conviertas esto en prosa.
4. No reformules.
5. No traduzcas autores/editoriales/ciudades/revistas/anos/numeros/paginas.
6. Solo anade traducciones al espanol entre parentesis para titulos claramente alemanes.
7. Si no estas seguro, no traduzcas.
8. No uses markdown.

Devuelve SOLO el texto final procesado.

TEXTO:
{text}
""".strip()
    resp = call_with_retries(
        lambda: client.chat.completions.create(
            model=cfg.model,
            temperature=cfg.temperature,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
        ),
        retries=cfg.retries,
        base_sleep_seconds=cfg.retry_base_sleep,
    )
    return (resp.choices[0].message.content or "").strip()


def _review_chunk(client: OpenAI, cfg: BibliographyConfig, source_text: str, draft_text: str) -> str:
    prompt = f"""
Revisa el bloque bibliografico ya procesado.

Verifica:
1. Numeracion y orden intactos.
2. Estructura y saltos de linea similares.
3. Solo traducciones para titulos alemanes.
4. Nada inventado.
5. Sin markdown.

Devuelve SOLO la version final corregida.

ORIGINAL:
{source_text}

PROPUESTA:
{draft_text}
""".strip()
    resp = call_with_retries(
        lambda: client.chat.completions.create(
            model=cfg.model,
            temperature=0.0,
            messages=[
                {"role": "system", "content": "Eres un revisor editorial especializado en bibliografia academica."},
                {"role": "user", "content": prompt},
            ],
        ),
        retries=cfg.retries,
        base_sleep_seconds=cfg.retry_base_sleep,
    )
    return (resp.choices[0].message.content or "").strip()


def run_bibliography_job(pdf_path: Path, output_dir: Path, cfg: BibliographyConfig | None = None) -> JobResult:
    cfg = cfg or BibliographyConfig()
    ensure_dir(output_dir)
    api_key = require_openai_api_key()
    client = OpenAI(api_key=api_key, timeout=cfg.timeout_seconds)

    pages = read_pdf_pages(pdf_path)
    selected_pages = slice_pages(pages, cfg.page_start, cfg.page_end)
    chunks = chunk_pages(selected_pages, cfg.pages_per_chunk)

    out_jsonl = output_dir / "bibliography_translated.jsonl"
    out_docx = output_dir / "bibliography_translated.docx"
    if out_jsonl.exists():
        out_jsonl.unlink()

    for chunk in chunks:
        source_text = truncate_text_if_needed(chunk.text, cfg.max_chars_per_chunk)
        draft = _process_chunk(client, cfg, source_text)
        reviewed = _review_chunk(client, cfg, chunk.text, draft)
        _append_jsonl(
            out_jsonl,
            {
                "chunk_id": chunk.chunk_id,
                "page_indices": chunk.page_indices,
                "source": chunk.text,
                "translation": reviewed,
            },
        )

    _build_docx_from_jsonl(out_jsonl, out_docx, add_headers=cfg.add_chunk_headers)
    return JobResult(
        tool="bibliography",
        status="done",
        message="Bibliography processing completed",
        artifacts=[
            Artifact(kind="jsonl", path=str(out_jsonl)),
            Artifact(kind="docx", path=str(out_docx)),
        ],
    )
