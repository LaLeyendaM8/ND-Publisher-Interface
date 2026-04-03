from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openai import OpenAI

from app.core.chunking import chunk_pages, slice_pages, truncate_text_if_needed
from app.core.config import require_openai_api_key
from app.core.files import ensure_dir, read_pdf_pages
from app.core.retry import call_with_retries
from app.core.schemas import Artifact, JobResult


SYSTEM_PROMPT = (
    "Actuas como traductor profesional especializado en textos filosoficos aleman-espanol. "
    "Traduces con precision conceptual, claridad academica y fidelidad al argumento del autor. "
    "Mantienes consistencia terminologica y estilo academico natural. "
    "No simplificas argumentos filosoficos complejos y respetas la estructura argumentativa del texto. "
    "Regla editorial de lenguaje: en espanol NO uses lenguaje inclusivo ni desdoblamientos. "
    "Usa siempre el masculino generico academico."
)


@dataclass
class TranslationConfig:
    pages_per_chunk: int = 3
    model: str = "gpt-4.1"
    temperature: float = 0.2
    timeout_seconds: float = 180.0
    retries: int = 5
    retry_base_sleep: float = 2.0
    max_chars_per_chunk: int = 12000
    source_lang: str = "Deutsch"
    target_lang: str = "Espanol"
    add_chunk_headers: bool = True
    page_start: int = 0
    page_end: int | None = None


def _load_glossary(path: Path | None) -> str:
    if not path or not path.exists():
        return ""
    return path.read_text(encoding="utf-8").strip()


def _parse_glossary_table(glossary_text: str) -> dict[str, str]:
    mapping: dict[str, str] = {}
    if not glossary_text:
        return mapping
    lines = [ln.strip() for ln in glossary_text.splitlines() if ln.strip()]
    for line in lines:
        if not line.startswith("|"):
            continue
        if "Aleman" in line or "Aleman" in line or "---" in line:
            continue
        parts = [p.strip() for p in line.strip("|").split("|")]
        if len(parts) < 2:
            continue
        de = parts[0]
        es = parts[1]
        if de and es:
            mapping[de] = es
    return mapping


def _append_jsonl(path: Path, row: dict[str, Any]) -> None:
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def _build_docx_from_jsonl(jsonl_path: Path, out_docx: Path, add_headers: bool) -> None:
    doc = Document()
    with jsonl_path.open("r", encoding="utf-8") as handle:
        for line in handle:
            obj = json.loads(line)
            if add_headers:
                pages = ", ".join(str(p + 1) for p in obj["page_indices"])
                doc.add_heading(f"Chunk {obj['chunk_id']} (PDF Seiten: {pages})", level=2)
            doc.add_paragraph(obj["translation"])
    doc.save(str(out_docx))


def _translate_chunk(client: OpenAI, cfg: TranslationConfig, text: str, glossary_text: str) -> str:
    glossary_map = _parse_glossary_table(glossary_text)
    glossary_block = ""
    if glossary_map:
        locked_lines = "\n".join([f"- {k} = {v}" for k, v in glossary_map.items()])
        glossary_block = (
            "\n\nGLOSARIO BLOQUEADO (OBLIGATORIO):\n"
            f"{locked_lines}\n"
            "Regla estricta: usa exactamente la traduccion bloqueada.\n"
        )

    user_prompt = f"""
Traduce el siguiente fragmento filosofico del {cfg.source_lang} al {cfg.target_lang}.

Requisitos:
- Fidelidad conceptual.
- Espanol academico natural.
- Manten la estructura argumentativa.
- Consistencia terminologica.

{glossary_block}

Devuelve SOLO la traduccion final en espanol, sin comentarios.

TEXTO (aleman):
{text}
""".strip()

    resp = call_with_retries(
        lambda: client.chat.completions.create(
            model=cfg.model,
            temperature=cfg.temperature,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
        ),
        retries=cfg.retries,
        base_sleep_seconds=cfg.retry_base_sleep,
    )
    return (resp.choices[0].message.content or "").strip()


def _review_translation(
    client: OpenAI,
    cfg: TranslationConfig,
    source_de: str,
    draft_es: str,
    glossary_text: str,
) -> str:
    glossary_map = _parse_glossary_table(glossary_text)
    locked_lines = "\n".join([f"- {k} = {v}" for k, v in glossary_map.items()]) if glossary_map else ""
    prompt = f"""
Revisa y mejora la siguiente traduccion filosofica (aleman -> espanol).

Objetivo:
- Corregir errores de sentido.
- Mejorar fluidez academica.
- Mantener consistencia terminologica.

GLOSARIO BLOQUEADO (OBLIGATORIO):
{locked_lines}

Devuelve SOLO la version final mejorada en espanol.

TEXTO ORIGINAL:
{source_de}

BORRADOR:
{draft_es}
""".strip()

    resp = call_with_retries(
        lambda: client.chat.completions.create(
            model=cfg.model,
            temperature=0.1,
            messages=[
                {"role": "system", "content": "Eres revisor editorial y traductor filosofico aleman-espanol."},
                {"role": "user", "content": prompt},
            ],
        ),
        retries=cfg.retries,
        base_sleep_seconds=cfg.retry_base_sleep,
    )
    return (resp.choices[0].message.content or "").strip()


def _enforce_glossary_terms(translation: str, glossary_text: str) -> str:
    corrected = translation
    glossary_map = _parse_glossary_table(glossary_text)
    for de, es in glossary_map.items():
        if de in corrected:
            corrected = corrected.replace(de, es)
    return corrected


def run_translation_job(
    pdf_path: Path,
    output_dir: Path,
    glossary_path: Path | None = None,
    cfg: TranslationConfig | None = None,
) -> JobResult:
    cfg = cfg or TranslationConfig()
    ensure_dir(output_dir)
    api_key = require_openai_api_key()
    client = OpenAI(api_key=api_key, timeout=cfg.timeout_seconds)

    pages = read_pdf_pages(pdf_path)
    selected_pages = slice_pages(pages, cfg.page_start, cfg.page_end)
    chunks = chunk_pages(selected_pages, cfg.pages_per_chunk)
    glossary_text = _load_glossary(glossary_path)

    out_jsonl = output_dir / "translation_chunks.jsonl"
    out_docx = output_dir / "translation.docx"
    if out_jsonl.exists():
        out_jsonl.unlink()

    for chunk in chunks:
        source_text = truncate_text_if_needed(chunk.text, cfg.max_chars_per_chunk)
        draft = _translate_chunk(client, cfg, source_text, glossary_text)
        reviewed = _review_translation(client, cfg, chunk.text, draft, glossary_text)
        final_translation = _enforce_glossary_terms(reviewed, glossary_text)
        _append_jsonl(
            out_jsonl,
            {
                "chunk_id": chunk.chunk_id,
                "page_indices": chunk.page_indices,
                "source_de": chunk.text,
                "translation": final_translation,
            },
        )

    _build_docx_from_jsonl(out_jsonl, out_docx, add_headers=cfg.add_chunk_headers)
    return JobResult(
        tool="translation",
        status="done",
        message="Translation completed",
        artifacts=[
            Artifact(kind="jsonl", path=str(out_jsonl)),
            Artifact(kind="docx", path=str(out_docx)),
        ],
    )
